import io
import json
from fpdf import FPDF
from docx import Document

def build_text_report(case_data: dict) -> str:
    am1 = case_data.get("am1", {})
    am2 = case_data.get("am2", {})
    am3 = case_data.get("am3", {})
    
    text = f"SMARTSNAKEBITE CLINICAL REPORT\n"
    text += f"Case ID: {case_data.get('id', 'N/A')}\n"
    text += f"Time: {case_data.get('time', 'N/A')}\n"
    text += f"Severity: {case_data.get('severity', 'UNKNOWN')}\n"
    text += "-" * 40 + "\n\n"
    
    text += "PATIENT TRANSCRIPT:\n"
    text += f"\"{case_data.get('transcript', '')}\"\n\n"
    
    text += "CLINICAL DECISION:\n"
    text += f"Referral Priority: {am2.get('referral_priority', 'N/A')}\n"
    text += f"Antivenom Required: {'Yes' if am2.get('antivenom_required') else 'No'}\n"
    text += f"Mortality Risk: {am2.get('mortality_risk', 'N/A')}\n"
    
    species = am2.get('top_species', am2.get('probable_species', []))
    if species:
        text += f"Probable Species: {', '.join(species)}\n"
    text += "\n"
    
    if am3.get("primary_corrective_message"):
        text += "CRITICAL SAFETY ALERT:\n"
        text += f"{am3['primary_corrective_message']}\n\n"
        
    explanation = am2.get("clinical_explanation", am2.get("explanation", ""))
    if explanation:
        if isinstance(explanation, list):
            explanation = "\n".join(explanation)
        text += "CLINICAL REASONING LOG:\n"
        text += f"{explanation}\n\n"
        
    text += "-" * 40 + "\n"
    text += "DISCLAIMER: This is an AI-assisted triage report. It is NOT a substitute for professional medical judgement.\n"
    
    return text

def build_pdf_report(case_data: dict) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", style="B", size=16)
    pdf.cell(200, 10, txt="SmartSnakebite Clinical Report", ln=1, align='C')
    
    pdf.set_font("Helvetica", size=10)
    pdf.cell(200, 10, txt=f"Case ID: {case_data.get('id', 'N/A')} | Time: {case_data.get('time', 'N/A')} | Severity: {case_data.get('severity', 'UNKNOWN')}", ln=1, align='C')
    pdf.ln(5)
    
    pdf.set_font("Helvetica", style="B", size=12)
    pdf.cell(200, 10, txt="Patient Transcript:", ln=1)
    pdf.set_font("Helvetica", style="I", size=11)
    pdf.multi_cell(0, 10, txt=f"\"{case_data.get('transcript', '')}\"")
    pdf.ln(5)
    
    am2 = case_data.get("am2", {})
    pdf.set_font("Helvetica", style="B", size=12)
    pdf.cell(200, 10, txt="Clinical Decision:", ln=1)
    pdf.set_font("Helvetica", size=11)
    
    pdf.cell(200, 8, txt=f"Referral Priority: {am2.get('referral_priority', 'N/A')}", ln=1)
    pdf.cell(200, 8, txt=f"Antivenom Required: {'Yes' if am2.get('antivenom_required') else 'No'}", ln=1)
    pdf.cell(200, 8, txt=f"Mortality Risk: {am2.get('mortality_risk', 'N/A')}", ln=1)
    
    species = am2.get('top_species', am2.get('probable_species', []))
    if species:
        pdf.cell(200, 8, txt=f"Probable Species: {', '.join(species)}", ln=1)
    pdf.ln(5)
    
    am3 = case_data.get("am3", {})
    if am3.get("primary_corrective_message"):
        pdf.set_font("Helvetica", style="B", size=12)
        pdf.set_text_color(220, 53, 69)
        pdf.cell(200, 10, txt="CRITICAL SAFETY ALERT:", ln=1)
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Helvetica", size=11)
        pdf.multi_cell(0, 10, txt=am3["primary_corrective_message"])
        pdf.ln(5)
        
    explanation = am2.get("clinical_explanation", am2.get("explanation", ""))
    if explanation:
        if isinstance(explanation, list):
            explanation = "\n".join(explanation)
        pdf.set_font("Helvetica", style="B", size=12)
        pdf.cell(200, 10, txt="Clinical Reasoning Log:", ln=1)
        pdf.set_font("Courier", size=9)
        # Handle unicode safely in FPDF1 by replacing non-latin or using basic formatting
        safe_expl = explanation.encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 5, txt=safe_expl)
        pdf.ln(5)
        
    pdf.set_font("Helvetica", style="I", size=8)
    pdf.set_text_color(100, 100, 100)
    pdf.multi_cell(0, 5, txt="DISCLAIMER: This is an AI-assisted triage report. It is NOT a substitute for professional medical judgement.")
    
    return pdf.output(dest='S')

def build_docx_report(case_data: dict) -> bytes:
    doc = Document()
    doc.add_heading('SmartSnakebite Clinical Report', 0)
    
    p = doc.add_paragraph()
    p.add_run(f"Case ID: {case_data.get('id', 'N/A')} | ").bold = True
    p.add_run(f"Severity: {case_data.get('severity', 'UNKNOWN')} | ").bold = True
    p.add_run(f"Time: {case_data.get('time', 'N/A')}")
    
    doc.add_heading('Patient Transcript', level=1)
    doc.add_paragraph(f"\"{case_data.get('transcript', '')}\"", style='Quote')
    
    am2 = case_data.get("am2", {})
    doc.add_heading('Clinical Decision', level=1)
    doc.add_paragraph(f"Referral Priority: {am2.get('referral_priority', 'N/A')}")
    doc.add_paragraph(f"Antivenom Required: {'Yes' if am2.get('antivenom_required') else 'No'}")
    doc.add_paragraph(f"Mortality Risk: {am2.get('mortality_risk', 'N/A')}")
    
    species = am2.get('top_species', am2.get('probable_species', []))
    if species:
        doc.add_paragraph(f"Probable Species: {', '.join(species)}")
        
    am3 = case_data.get("am3", {})
    if am3.get("primary_corrective_message"):
        doc.add_heading('CRITICAL SAFETY ALERT', level=1)
        doc.add_paragraph(am3["primary_corrective_message"])
        
    explanation = am2.get("clinical_explanation", am2.get("explanation", ""))
    if explanation:
        if isinstance(explanation, list):
            explanation = "\n".join(explanation)
        doc.add_heading('Clinical Reasoning Log', level=1)
        doc.add_paragraph(explanation)
        
    doc.add_paragraph("DISCLAIMER: This is an AI-assisted triage report. It is NOT a substitute for professional medical judgement.")
    
    f = io.BytesIO()
    doc.save(f)
    return f.getvalue()
